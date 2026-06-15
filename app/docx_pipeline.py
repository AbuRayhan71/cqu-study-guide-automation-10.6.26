from __future__ import annotations

import copy
import re
import shutil
import tempfile
import zipfile
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Iterable, Literal

from lxml import etree
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph


BlockType = Literal["heading", "paragraph", "bullet_list", "numbered_list", "table", "image", "quote"]


@dataclass
class GenerationMetadata:
    unit_code: str = ""
    unit_name: str = ""
    week_number: str = ""
    week_title: str = ""
    version: str = ""
    year: str = ""
    term: str = ""


@dataclass
class ContentBlock:
    type: BlockType
    text: str = ""
    level: int = 1
    number: str = ""
    items: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)
    image_id: str = ""
    caption: str = ""


@dataclass
class ExtractedImage:
    image_id: str
    blob: bytes
    filename: str
    content_type: str


@dataclass
class StructuredDocument:
    metadata: GenerationMetadata
    body_blocks: list[ContentBlock]
    images: dict[str, ExtractedImage]


@dataclass
class GenerationResult:
    output_path: Path
    block_count: int
    image_count: int
    table_count: int


NAMESPACES = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
}

DOCUMENT_RELATIONSHIP_TYPES = {
    "footnotes.xml": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
    "endnotes.xml": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes",
}


def iter_block_items(parent: DocumentObject) -> Iterable[Paragraph | Table]:
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def extract_docx(path: Path, overrides: GenerationMetadata | None = None) -> StructuredDocument:
    document = Document(path)
    blocks: list[ContentBlock] = []
    images: dict[str, ExtractedImage] = {}
    image_counter = 1

    for item in iter_block_items(document):
        if isinstance(item, Paragraph):
            text = clean_text(item.text)
            style_name = item.style.name if item.style else ""
            for rid in image_relationship_ids(item):
                image_id = f"img_{image_counter:03d}"
                image_counter += 1
                related = document.part.related_parts[rid]
                filename = Path(getattr(related, "partname", f"{image_id}.bin")).name
                images[image_id] = ExtractedImage(
                    image_id=image_id,
                    blob=related.blob,
                    filename=filename,
                    content_type=getattr(related, "content_type", "application/octet-stream"),
                )
                blocks.append(ContentBlock(type="image", image_id=image_id))
            if not text:
                continue
            if style_name.lower().startswith("heading"):
                blocks.append(ContentBlock(type="heading", level=heading_level(style_name), text=text))
            elif is_numbered_paragraph(item):
                blocks.append(ContentBlock(type="numbered_list", items=[text]))
            elif is_bullet_paragraph(item):
                blocks.append(ContentBlock(type="bullet_list", items=[text]))
            else:
                blocks.append(ContentBlock(type="paragraph", text=text))
        else:
            rows = [[clean_text(cell.text) for cell in row.cells] for row in item.rows]
            if rows:
                blocks.append(ContentBlock(type="table", rows=rows))

    metadata = infer_metadata(blocks)
    if overrides:
        metadata = merge_metadata(metadata, overrides)
    return StructuredDocument(metadata=metadata, body_blocks=coalesce_lists(blocks), images=images)


def generate_study_guide(
    input_docx: Path,
    template_docx: Path,
    output_docx: Path,
    metadata: GenerationMetadata | None = None,
) -> GenerationResult:
    structured = extract_docx(input_docx, metadata)
    render_preserving_source_document(input_docx, template_docx, output_docx, structured.metadata)
    validate_output(output_docx, structured)
    return GenerationResult(
        output_path=output_docx,
        block_count=len(structured.body_blocks),
        image_count=len(structured.images),
        table_count=sum(1 for block in structured.body_blocks if block.type == "table"),
    )


def render_preserving_source_document(
    source_path: Path,
    template_path: Path,
    output_path: Path,
    metadata: GenerationMetadata,
    accepted_corrections: list[dict[str, str]] | None = None,
) -> None:
    """Render by keeping source OOXML body content intact.

    This is the default path for real study guides because it preserves the
    source Contents section, heading styles, numbering, tables, fields, and
    embedded images much better than reconstructing from extracted plain text.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        rendered_template = tmp_path / "template_rendered.docx"
        base_dir = tmp_path / "base"
        source_dir = tmp_path / "source"

        template_doc = Document(template_path)
        replace_placeholders(template_doc, metadata, remove_toc_placeholder=True)
        template_doc.save(rendered_template)

        unzip_docx(rendered_template, base_dir)
        unzip_docx(source_path, source_dir)

        copy_package_parts(source_dir, base_dir)
        rel_map = merge_document_relationships(source_dir, base_dir)
        replace_template_body_with_source(base_dir, source_dir, accepted_corrections or [], rel_map)
        prune_unused_styles(base_dir)
        zip_docx(base_dir, output_path)
        Document(output_path)


def clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def image_relationship_ids(paragraph: Paragraph) -> list[str]:
    ids: list[str] = []
    for blip in paragraph._element.xpath(".//a:blip"):
        rid = blip.get(qn("r:embed"))
        if rid:
            ids.append(rid)
    return ids


def heading_level(style_name: str) -> int:
    match = re.search(r"(\d+)", style_name)
    if not match:
        return 1
    return max(1, min(9, int(match.group(1))))


def is_bullet_paragraph(paragraph: Paragraph) -> bool:
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    return "bullet" in style_name


def is_numbered_paragraph(paragraph: Paragraph) -> bool:
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    return "number" in style_name or "list paragraph" in style_name


def coalesce_lists(blocks: list[ContentBlock]) -> list[ContentBlock]:
    coalesced: list[ContentBlock] = []
    for block in blocks:
        if block.type in {"bullet_list", "numbered_list"} and coalesced and coalesced[-1].type == block.type:
            coalesced[-1].items.extend(block.items)
        else:
            coalesced.append(block)
    return coalesced


def infer_metadata(blocks: list[ContentBlock]) -> GenerationMetadata:
    joined = "\n".join(block.text for block in blocks if block.text)
    unit_code = first_match(joined, r"\b[A-Z]{4}\d{5}\b")
    week_number = first_match(joined, r"\bWEEK\s+(\d{1,2})\b", group=1)
    unit_name = ""
    for block in blocks[:10]:
        if block.text and block.text != unit_code and not re.search(r"\bSTUDY GUIDE\b", block.text, re.I):
            unit_name = block.text
            break
    return GenerationMetadata(unit_code=unit_code, unit_name=unit_name, week_number=week_number)


def first_match(text: str, pattern: str, group: int = 0) -> str:
    match = re.search(pattern, text, re.I)
    return match.group(group).strip() if match else ""


def merge_metadata(base: GenerationMetadata, override: GenerationMetadata) -> GenerationMetadata:
    return GenerationMetadata(
        unit_code=override.unit_code or base.unit_code,
        unit_name=override.unit_name or base.unit_name,
        week_number=override.week_number or base.week_number,
        week_title=override.week_title or base.week_title,
        version=override.version or base.version,
        year=override.year or base.year,
        term=override.term or base.term,
    )


def render_document(structured: StructuredDocument, template_path: Path, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document(template_path)
    replace_placeholders(document, structured.metadata)
    body_placeholder = find_paragraph_containing(document, "{{ body_subdoc }}")
    if body_placeholder is None:
        body_placeholder = document.add_paragraph()
    clear_paragraph(body_placeholder)

    anchor = body_placeholder
    for block in structured.body_blocks:
        anchor = insert_block_after(anchor, block, document, structured.images)

    save_and_reload(document, output_path)


def replace_placeholders(
    document: DocumentObject,
    metadata: GenerationMetadata,
    remove_toc_placeholder: bool = False,
) -> None:
    replacements = {
        "{{ unit_code }}": metadata.unit_code or "UNIT CODE",
        "{{ unit_name | upper }}": (metadata.unit_name or "UNIT NAME").upper(),
        "STUDY GUIDE | WEEK {{ week_number }}": f"STUDY GUIDE | WEEK {metadata.week_number or '1'}",
        "{{ week_title }}": metadata.week_title,
        "{{ version }}": metadata.version,
        "{{ toc_placeholder }}": "" if remove_toc_placeholder else "Update this table of contents in Word after opening the generated document.",
    }
    for paragraph in document.paragraphs:
        text = paragraph.text
        normalized = text.replace("{{ toc _placeholder }}", "{{ toc_placeholder }}")
        normalized = normalized.replace("{{ body _subdoc }}", "{{ body_subdoc }}")
        if normalized != text:
            set_paragraph_text(paragraph, normalized)
            text = normalized
        for placeholder, replacement in replacements.items():
            if placeholder in text:
                set_paragraph_text(paragraph, text.replace(placeholder, replacement))
                text = paragraph.text


def find_paragraph_containing(document: DocumentObject, needle: str) -> Paragraph | None:
    for paragraph in document.paragraphs:
        normalized = paragraph.text.replace("{{ body _subdoc }}", "{{ body_subdoc }}")
        if needle in normalized:
            return paragraph
    return None


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    clear_paragraph(paragraph)
    paragraph.add_run(text)


def insert_block_after(
    anchor: Paragraph,
    block: ContentBlock,
    document: DocumentObject,
    images: dict[str, ExtractedImage],
) -> Paragraph:
    if block.type == "heading":
        paragraph = insert_paragraph_after(anchor, block.text, style=f"Heading {block.level}")
        return paragraph
    if block.type == "paragraph":
        return insert_paragraph_after(anchor, block.text)
    if block.type == "quote":
        paragraph = insert_paragraph_after(anchor, block.text)
        paragraph.style = "Quote" if "Quote" in [style.name for style in document.styles] else paragraph.style
        if block.citation:
            return insert_paragraph_after(paragraph, block.citation)
        return paragraph
    if block.type in {"bullet_list", "numbered_list"}:
        current = anchor
        style = "List Bullet" if block.type == "bullet_list" else "List Number"
        for item in block.items:
            current = insert_paragraph_after(current, item, style=style)
        return current
    if block.type == "table":
        table = document.add_table(rows=0, cols=max(len(row) for row in block.rows))
        table.style = "Table Grid"
        for row_values in block.rows:
            cells = table.add_row().cells
            for idx, value in enumerate(row_values):
                cells[idx].text = value
        anchor._p.addnext(table._tbl)
        marker = insert_paragraph_after(anchor, "")
        table._tbl.addnext(marker._p)
        return marker
    if block.type == "image":
        paragraph = insert_paragraph_after(anchor, "")
        image = images.get(block.image_id)
        if image is None:
            raise ValueError(f"Missing extracted image {block.image_id}")
        run = paragraph.add_run()
        run.add_picture(BytesIO(image.blob), width=Inches(5.8))
        if block.caption:
            return insert_paragraph_after(paragraph, block.caption)
        return paragraph
    return anchor


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = copy.deepcopy(paragraph._p)
    for child in list(new_p):
        new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_paragraph.style = style
        except KeyError:
            pass
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def save_and_reload(document: DocumentObject, output_path: Path) -> None:
    document.save(output_path)
    Document(output_path)


def validate_output(path: Path, structured: StructuredDocument) -> None:
    if not path.exists() or path.stat().st_size == 0:
        raise ValueError("Output DOCX was not created.")
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    unresolved = re.findall(r"{{[^}]+}}|{%[^%]+%}", text)
    if unresolved:
        raise ValueError(f"Unresolved template placeholders remain: {', '.join(unresolved)}")
    if any(block.type == "heading" for block in structured.body_blocks) and not any(
        paragraph.style and paragraph.style.name.lower().startswith("heading") for paragraph in document.paragraphs
    ):
        raise ValueError("Expected headings were not present in output.")
    expected_tables = sum(1 for block in structured.body_blocks if block.type == "table")
    if expected_tables and len(document.tables) == 0 and "CONTENTS" not in text.upper():
        raise ValueError("One or more source tables were not present in output.")
    expected_images = len(structured.images)
    actual_images = count_docx_media(path)
    template_images = 4
    if expected_images and actual_images < template_images + expected_images:
        raise ValueError("One or more source images were not present in output.")
    missing_notes = missing_note_references(path)
    if missing_notes:
        raise ValueError(f"Missing note definitions in output: {', '.join(missing_notes)}")
    style_issues = invalid_style_references(path)
    if style_issues:
        raise ValueError(f"Invalid style references in output: {', '.join(style_issues)}")


def count_docx_media(path: Path) -> int:
    with zipfile.ZipFile(path) as archive:
        return len([name for name in archive.namelist() if name.startswith("word/media/")])


def missing_note_references(path: Path) -> list[str]:
    missing: list[str] = []
    with zipfile.ZipFile(path) as archive:
        document_root = etree.fromstring(archive.read("word/document.xml"))
        for note_type, part_name in (("footnote", "word/footnotes.xml"), ("endnote", "word/endnotes.xml")):
            references = {
                node.get(qn("w:id"))
                for node in document_root.xpath(f".//w:{note_type}Reference", namespaces=NAMESPACES)
            }
            references.discard(None)
            if not references:
                continue
            if part_name not in archive.namelist():
                missing.extend(f"{note_type}:{note_id}" for note_id in sorted(references))
                continue
            notes_root = etree.fromstring(archive.read(part_name))
            definitions = {
                node.get(qn("w:id"))
                for node in notes_root.xpath(f".//w:{note_type}", namespaces=NAMESPACES)
            }
            missing.extend(f"{note_type}:{note_id}" for note_id in sorted(references - definitions))
    return missing


def invalid_style_references(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as archive:
        styles_root = etree.fromstring(archive.read("word/styles.xml"))
        defined = {
            style.get(qn("w:styleId"))
            for style in styles_root.findall("w:style", namespaces=NAMESPACES)
            if style.get(qn("w:styleId"))
        }
        issues = set()
        for name in archive.namelist():
            if not name.startswith("word/") or not name.endswith(".xml") or name == "word/styles.xml":
                continue
            root = etree.fromstring(archive.read(name))
            for tag in ("w:pStyle", "w:rStyle", "w:tblStyle"):
                for node in root.findall(f".//{tag}", namespaces=NAMESPACES):
                    style_id = node.get(qn("w:val"))
                    if style_id and style_id not in defined:
                        issues.add(f"{name}:{style_id}")
        for tag in ("w:basedOn", "w:link", "w:next"):
            for node in styles_root.findall(f".//{tag}", namespaces=NAMESPACES):
                style_id = node.get(qn("w:val"))
                if style_id and style_id not in defined:
                    issues.add(f"styles.xml:{style_id}")
        return sorted(issues)


def unzip_docx(path: Path, destination: Path) -> None:
    destination.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path) as archive:
        archive.extractall(destination)


def zip_docx(source_dir: Path, output_path: Path) -> None:
    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for path in source_dir.rglob("*"):
            if path.is_file():
                archive.write(path, path.relative_to(source_dir).as_posix())


def parse_xml(path: Path) -> etree._ElementTree:
    return etree.parse(str(path))


def write_xml(tree: etree._ElementTree, path: Path) -> None:
    tree.write(str(path), xml_declaration=True, encoding="UTF-8", standalone=True)


def copy_package_parts(source_dir: Path, base_dir: Path) -> None:
    copy_if_exists(source_dir / "word" / "numbering.xml", base_dir / "word" / "numbering.xml")
    copy_word_part_with_relationships(source_dir, base_dir, "footnotes.xml")
    copy_word_part_with_relationships(source_dir, base_dir, "endnotes.xml")
    merge_missing_styles(source_dir / "word" / "styles.xml", base_dir / "word" / "styles.xml")
    merge_content_types(source_dir / "[Content_Types].xml", base_dir / "[Content_Types].xml")


def copy_tree(source: Path, destination: Path) -> None:
    if not source.exists():
        return
    destination.mkdir(parents=True, exist_ok=True)
    for path in source.rglob("*"):
        if path.is_file():
            target = destination / path.relative_to(source)
            target.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(path, target)


def copy_if_exists(source: Path, destination: Path) -> None:
    if source.exists():
        destination.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, destination)


def copy_word_part_with_relationships(source_dir: Path, base_dir: Path, part_name: str) -> None:
    source_part = source_dir / "word" / part_name
    if not source_part.exists():
        return
    copy_if_exists(source_part, base_dir / "word" / part_name)
    relationship_type = DOCUMENT_RELATIONSHIP_TYPES.get(part_name)
    if relationship_type:
        ensure_document_relationship(base_dir, relationship_type, part_name)

    source_rels = source_dir / "word" / "_rels" / f"{part_name}.rels"
    if not source_rels.exists():
        return

    rels_tree = parse_xml(source_rels)
    for rel in rels_tree.getroot():
        target = rel.get("Target")
        mode = rel.get("TargetMode")
        if target and mode != "External":
            rel.set("Target", copy_related_part(source_dir, base_dir, target))
    write_xml(rels_tree, base_dir / "word" / "_rels" / f"{part_name}.rels")


def ensure_document_relationship(base_dir: Path, relationship_type: str, target: str) -> None:
    rels_path = base_dir / "word" / "_rels" / "document.xml.rels"
    rels_tree = parse_xml(rels_path)
    root = rels_tree.getroot()
    for rel in root:
        if rel.get("Type") == relationship_type:
            rel.set("Target", target)
            write_xml(rels_tree, rels_path)
            return
    existing_ids = {rel.get("Id") for rel in root}
    relationship = etree.SubElement(root, f"{{{NAMESPACES['rel']}}}Relationship")
    relationship.set("Id", f"rId{next_relationship_number(existing_ids)}")
    relationship.set("Type", relationship_type)
    relationship.set("Target", target)
    write_xml(rels_tree, rels_path)


def merge_missing_styles(source_styles: Path, base_styles: Path) -> None:
    if not source_styles.exists() or not base_styles.exists():
        return
    source_tree = parse_xml(source_styles)
    base_tree = parse_xml(base_styles)
    base_root = base_tree.getroot()
    source_root = source_tree.getroot()
    existing = {style.get(qn("w:styleId")): style for style in base_root.findall("w:style", namespaces=NAMESPACES)}
    for style in source_root.findall("w:style", namespaces=NAMESPACES):
        style_id = style.get(qn("w:styleId"))
        if not style_id:
            continue
        if style_id not in existing:
            base_root.append(copy.deepcopy(style))
            existing[style_id] = style
    write_xml(base_tree, base_styles)


def prune_unused_styles(base_dir: Path) -> None:
    styles_path = base_dir / "word" / "styles.xml"
    if not styles_path.exists():
        return
    styles_tree = parse_xml(styles_path)
    styles_root = styles_tree.getroot()
    used_style_ids = collect_used_style_ids(base_dir)
    style_nodes = {
        style.get(qn("w:styleId")): style
        for style in styles_root.findall("w:style", namespaces=NAMESPACES)
        if style.get(qn("w:styleId"))
    }
    for style_id, style in style_nodes.items():
        if is_builtin_template_style(style):
            used_style_ids.add(style_id)

    changed = True
    while changed:
        changed = False
        for style_id in list(used_style_ids):
            style = style_nodes.get(style_id)
            if style is None:
                continue
            for tag in ("w:basedOn", "w:link", "w:next"):
                related = style.find(tag, namespaces=NAMESPACES)
                related_id = related.get(qn("w:val")) if related is not None else None
                if related_id and related_id not in used_style_ids:
                    used_style_ids.add(related_id)
                    changed = True

    for style_id, style in list(style_nodes.items()):
        if style_id in used_style_ids:
            continue
        styles_root.remove(style)
    write_xml(styles_tree, styles_path)


def collect_used_style_ids(base_dir: Path) -> set[str]:
    used: set[str] = set()
    for path in (base_dir / "word").glob("*.xml"):
        if path.name == "styles.xml":
            continue
        root = parse_xml(path).getroot()
        for tag in ("w:pStyle", "w:rStyle", "w:tblStyle"):
            for node in root.findall(f".//{tag}", namespaces=NAMESPACES):
                style_id = node.get(qn("w:val"))
                if style_id:
                    used.add(style_id)
    return used


def is_builtin_template_style(style: etree._Element) -> bool:
    style_id = style.get(qn("w:styleId")) or ""
    custom = style.get(qn("w:customStyle"))
    name = style.find("w:name", namespaces=NAMESPACES)
    if custom == "1":
        return False
    if name is None:
        return False
    return bool(style_id)


def merge_content_types(source_types: Path, base_types: Path) -> None:
    if not source_types.exists() or not base_types.exists():
        return
    source_tree = parse_xml(source_types)
    base_tree = parse_xml(base_types)
    base_root = base_tree.getroot()
    existing_defaults = {
        item.get("Extension")
        for item in base_root.findall("{http://schemas.openxmlformats.org/package/2006/content-types}Default")
    }
    existing_overrides = {
        item.get("PartName")
        for item in base_root.findall("{http://schemas.openxmlformats.org/package/2006/content-types}Override")
    }
    for item in source_tree.getroot():
        extension = item.get("Extension")
        part_name = item.get("PartName")
        if extension and extension not in existing_defaults:
            base_root.append(copy.deepcopy(item))
            existing_defaults.add(extension)
        elif part_name and part_name not in existing_overrides:
            base_root.append(copy.deepcopy(item))
            existing_overrides.add(part_name)
    write_xml(base_tree, base_types)


def merge_document_relationships(source_dir: Path, base_dir: Path) -> dict[str, str]:
    source_rels = source_dir / "word" / "_rels" / "document.xml.rels"
    base_rels = base_dir / "word" / "_rels" / "document.xml.rels"
    if not source_rels.exists():
        return {}
    source_tree = parse_xml(source_rels)
    base_tree = parse_xml(base_rels)
    base_root = base_tree.getroot()
    existing_ids = {rel.get("Id") for rel in base_root}
    next_id = next_relationship_number(existing_ids)
    rel_map: dict[str, str] = {}
    for rel in source_tree.getroot():
        old_id = rel.get("Id")
        if not old_id:
            continue
        if not should_merge_relationship(rel.get("Type", "")):
            continue
        new_id = old_id
        if new_id in existing_ids:
            new_id = f"rId{next_id}"
            next_id += 1
        new_rel = copy.deepcopy(rel)
        new_rel.set("Id", new_id)
        target = rel.get("Target")
        mode = rel.get("TargetMode")
        if target and mode != "External":
            new_rel.set("Target", copy_related_part(source_dir, base_dir, target))
        base_root.append(new_rel)
        existing_ids.add(new_id)
        rel_map[old_id] = new_id
    write_xml(base_tree, base_rels)
    return rel_map


def should_merge_relationship(rel_type: str) -> bool:
    body_relationships = (
        "/image",
        "/hyperlink",
        "/chart",
        "/oleObject",
        "/package",
        "/diagram",
        "/audio",
        "/video",
    )
    return any(rel_type.endswith(suffix) for suffix in body_relationships)


def next_relationship_number(ids: set[str | None]) -> int:
    numbers = []
    for rel_id in ids:
        match = re.fullmatch(r"rId(\d+)", rel_id or "")
        if match:
            numbers.append(int(match.group(1)))
    return max(numbers, default=0) + 1


def replace_template_body_with_source(
    base_dir: Path,
    source_dir: Path,
    accepted_corrections: list[dict[str, str]],
    rel_map: dict[str, str],
) -> None:
    base_doc = base_dir / "word" / "document.xml"
    source_doc = source_dir / "word" / "document.xml"
    base_tree = parse_xml(base_doc)
    source_tree = parse_xml(source_doc)
    base_body = base_tree.find(".//w:body", namespaces=NAMESPACES)
    source_body = source_tree.find(".//w:body", namespaces=NAMESPACES)
    if base_body is None or source_body is None:
        raise ValueError("DOCX body could not be read.")

    source_elements = trim_source_front_matter([
        copy.deepcopy(child)
        for child in source_body
        if child.tag != qn("w:sectPr")
    ])
    for element in source_elements:
        apply_corrections_to_xml_element(element, accepted_corrections)
        normalize_source_body_formatting(element)
        remap_relationship_ids(element, rel_map)

    insert_index = find_body_placeholder_index(base_body)
    if insert_index is None:
        insert_index = len(base_body)
    else:
        remove_template_toc_region(base_body, insert_index)
        insert_index = find_body_placeholder_index(base_body)
        if insert_index is None:
            insert_index = len(base_body)
        else:
            del base_body[insert_index]

    for offset, element in enumerate(source_elements):
        base_body.insert(insert_index + offset, element)
    write_xml(base_tree, base_doc)


def find_body_placeholder_index(body: etree._Element) -> int | None:
    for index, child in enumerate(body):
        if "{{ body_subdoc }}" in element_text(child).replace("{{ body _subdoc }}", "{{ body_subdoc }}"):
            return index
    return None


def remove_template_toc_region(body: etree._Element, body_placeholder_index: int) -> None:
    start = None
    for index, child in enumerate(body[: body_placeholder_index + 1]):
        text = element_text(child).strip().upper()
        if text == "CONTENTS":
            start = index
            break
    if start is None:
        return
    for index in range(body_placeholder_index - 1, start - 1, -1):
        del body[index]


def element_text(element: etree._Element) -> str:
    return "".join(element.xpath(".//w:t/text()", namespaces=NAMESPACES))


def trim_source_front_matter(elements: list[etree._Element]) -> list[etree._Element]:
    for index, element in enumerate(elements):
        if element_text(element).strip().upper() == "CONTENTS":
            return elements[index:]
    return elements


def normalize_source_body_formatting(element: etree._Element) -> None:
    """Keep source structure, but let the CQU template control typography."""
    for run_properties in element.xpath(".//w:rPr", namespaces=NAMESPACES):
        parent = run_properties.getparent()
        if parent is not None:
            parent.remove(run_properties)

    for paragraph_properties in element.xpath(".//w:pPr", namespaces=NAMESPACES):
        for child in list(paragraph_properties):
            if child.tag not in {qn("w:pStyle"), qn("w:numPr"), qn("w:tabs"), qn("w:pageBreakBefore")}:
                paragraph_properties.remove(child)

    for table_properties in element.xpath(".//w:tblPr", namespaces=NAMESPACES):
        for child in list(table_properties):
            if child.tag != qn("w:tblStyle"):
                table_properties.remove(child)


def apply_corrections_to_xml_element(element: etree._Element, corrections: list[dict[str, str]]) -> None:
    if not corrections:
        return
    text_nodes = element.xpath(".//w:t", namespaces=NAMESPACES)
    if not text_nodes:
        return
    combined = "".join(node.text or "" for node in text_nodes)
    updated = combined
    for correction in corrections:
        original = correction.get("original", "")
        replacement = correction.get("replacement", "")
        if original and replacement and original in updated:
            updated = updated.replace(original, replacement, 1)
    if updated == combined:
        return
    text_nodes[0].text = updated
    for node in text_nodes[1:]:
        node.text = ""


def copy_related_part(source_dir: Path, base_dir: Path, target: str) -> str:
    if target.startswith("/") or ".." in Path(target).parts:
        return target
    source_file = source_dir / "word" / target
    if not source_file.exists():
        return target
    destination = base_dir / "word" / target
    if destination.exists():
        destination = unique_related_destination(destination)
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source_file, destination)
    return destination.relative_to(base_dir / "word").as_posix()


def unique_related_destination(destination: Path) -> Path:
    stem = destination.stem
    suffix = destination.suffix
    parent = destination.parent
    counter = 1
    while True:
        candidate = parent / f"{stem}_source_{counter}{suffix}"
        if not candidate.exists():
            return candidate
        counter += 1


def remap_relationship_ids(element: etree._Element, rel_map: dict[str, str]) -> None:
    if not rel_map:
        return
    for node in element.iter():
        for attr_name, attr_value in list(node.attrib.items()):
            if attr_name.startswith(f"{{{NAMESPACES['r']}}}") and attr_value in rel_map:
                node.set(attr_name, rel_map[attr_value])
